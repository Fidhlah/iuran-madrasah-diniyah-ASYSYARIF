import docx
from docx.shared import Pt

path = r'C:\Users\Fidh\AppData\Local\hermes\tmp\Laporan Keuangan Bulan Juli 2026 (dgn KOP).docx'
doc = docx.Document(path)

print("=== PARAGRAPHS (non-kosong) ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        # deteksi alignment
        al = str(p.alignment)
        runs = []
        for r in p.runs:
            if r.text.strip():
                runs.append(f"({r.text[:30]}|b={r.font.bold})")
        print(f"  P{i} [{al}] {t[:80]}")

print(f"\n=== JUMLAH TABEL: {len(doc.tables)} ===")
for ti, t in enumerate(doc.tables):
    print(f"\n--- TABEL {ti} ({len(t.rows)} rows x {len(t.columns)} cols) ---")
    for ri, row in enumerate(t.rows):
        cells = [c.text.strip()[:35] for c in row.cells]
        print(f"  R{ri}: {cells}")