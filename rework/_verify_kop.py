import zipfile
import docx
from xml.etree import ElementTree as ET

path = r'C:\Users\Fidh\AppData\Local\hermes\tmp\Laporan Keuangan Bulan Juli 2026 (dgn KOP).docx'

# 1. Cek kop masih ada (header image)
with zipfile.ZipFile(path) as z:
    has_media = any('media/image1.png' in n for n in z.namelist())
    has_header = any('header1.xml' in n for n in z.namelist())
    # cek isi body
    xml = z.read('word/document.xml').decode('utf-8')
    body_text = ''.join(t.text or '' for t in ET.fromstring(xml).iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
print(f"KOP: media={has_media}, header={has_header}")

doc = docx.Document(path)
print(f"Tabel: {len(doc.tables)} (harusnya 1 ringkasan + 1 lampiran)")
print(f"\n=== RINGKASAN ===")
for row in doc.tables[0].rows:
    print(f"  {[c.text.strip() for c in row.cells]}")

lamp = doc.tables[1]
print(f"\n=== LAMPIRAN: {len(lamp.rows)-1} transaksi ===")
for ri in [0,1,2,66,67,68]:
    print(f"  R{ri}: {[c.text.strip()[:30] for c in lamp.rows[ri].cells]}")