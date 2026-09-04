import docx

path = r'C:\Users\Fidh\AppData\Local\hermes\tmp\Laporan Keuangan Bulan Juli 2026 (dgn KOP).docx'
doc = docx.Document(path)
ring = doc.tables[0]
lamp = doc.tables[1]

# Ambil nilai ringkasan dengan index baris eksplisit (bukan by label — labels duplikat)
def ring_idx(i):
    return ring.cell(i, 1).text.strip()

print("=== TABEL RINGKASAN (dgn index) ===")
for i, row in enumerate(ring.rows):
    print(f"  R{i}: {row.cells[0].text.strip():25} {row.cells[1].text.strip()}")

# ── Hitung ulang dari Lampiran ──
income_total = expense_total = 0
income_count = expense_count = 0
spp_bulan = spp_tunggakan = pendaftaran = 0
rows_data = []
for r in lamp.rows[1:]:
    cells = [c.text.strip() for c in r.cells]
    if len(cells) < 5 or not cells[0].isdigit():
        continue
    no, tgl, jenis, ket, jml = cells[0], cells[1], cells[2], cells[3], cells[4]
    amt = int(jml)
    rows_data.append((no, tgl, jenis, ket, amt))
    if jenis == "Pemasukan":
        income_total += amt
        income_count += 1
        if "pendaftaran" in ket.lower():
            pendaftaran += amt
        elif "bulan juli" in ket.lower():
            spp_bulan += amt
        else:
            spp_tunggakan += amt
    else:
        expense_total += amt
        expense_count += 1

print(f"\nLampiran: {len(rows_data)} tx = {income_count} pemasukan ({income_total:,}) + {expense_count} pengeluaran ({expense_total:,})")

# Saldo (pakai index eksplisit)
# R0 SALDO AWAL, R1 KB awal, R2 MDTA awal ; R11 KB akhir, R12 MDTA akhir, R13 total akhir
kb_awal = int(ring_idx(1).replace("Rp ","").replace(".",""))
mdta_awal = int(ring_idx(2).replace("Rp ","").replace(".",""))
kb_akhir = int(ring_idx(11).replace("Rp ","").replace(".",""))
mdta_akhir = int(ring_idx(12).replace("Rp ","").replace(".",""))
total_akhir = int(ring_idx(13).replace("Rp ","").replace(".",""))
spp_file = int(ring_idx(4).replace("Rp ","").replace(".",""))
tungg_file = int(ring_idx(5).replace("Rp ","").replace(".",""))
pendaft_file = int(ring_idx(6).replace("Rp ","").replace(".",""))
total_masuk_file = int(ring_idx(7).replace("Rp ","").replace(".",""))
riil_file = int(ring_idx(8).replace("Rp ","").replace(".",""))
alokasi_file = int(ring_idx(9).replace("Rp ","").replace(".",""))

def cek(label, a, b):
    print(f"  {label:42} a={a:,} b={b:,} → {'OK' if a==b else '❌ MISMATCH'}")

print("\n=== VERIFIKASI AKURAT ===")
cek("SPP Bulan Ini (file vs lampiran)", spp_file, spp_bulan)
cek("SPP Tunggakan (file vs lampiran)", tungg_file, spp_tunggakan)
cek("Uang Pendaftaran (file vs lampiran)", pendaft_file, pendaftaran)
cek("Total Pemasukan (file vs lampiran)", total_masuk_file, income_total)
cek("Pengeluaran RIIL (file vs expense-alokasi)", riil_file, expense_total - alokasi_file)
cek("Saldo Awal total (KB+MDTA)", kb_awal + mdta_awal, 4881000)
cek("Saldo Akhir KB (file)", kb_akhir, 3950000)
cek("Saldo Akhir MDTA (file)", mdta_akhir, 1915100)
cek("Total Akhir (file)", total_akhir, 5865100)

# Logika saldo akhir
print("\n=== CEK SILANG SALDO ===")
kb_calc = kb_awal + income_total - (expense_total - alokasi_file)
exc_gaji_mesjid = sum(a for _,_,j,k,a in rows_data if j=="Pengeluaran" and (("honor" in k.lower()) or ("mesjid" in k.lower())))
kb_calc2 = kb_awal + (spp_bulan + spp_tunggakan + pendaftaran) - exc_gaji_mesjid - alokasi_file
print(f"  KB akhir rumus: {kb_awal:,} + {income_total:,} - ({expense_total:,}-{alokasi_file:,}) = {kb_calc:,} (file {kb_akhir:,})")
mdta_calc = mdta_awal + alokasi_file - sum(a for _,_,j,k,a in rows_data if j=="Pengeluaran" and "kas mdta" not in k.lower() and "honor" not in k.lower() and "mesjid" not in k.lower())
print(f"  MDTA akhir rumus: {mdta_awal:,} + {alokasi_file:,} - belanja = {mdta_calc:,} (file {mdta_akhir:,})")

# Kas MDTA & urutan
mdta = [r for r in rows_data if 'kas mdta' in r[3].lower()]
print(f"\nKas MDTA di Lampiran: {mdta}")