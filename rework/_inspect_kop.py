import docx
from docx.shared import Pt

path = r'D:\fidh\Asysyarif\Dokumen\KOP-MDTA.docx'
doc = docx.Document(path)

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    runs_info = []
    for r in p.runs:
        runs_info.append(f"[{r.font.name} {r.font.size} bold={r.font.bold}] '{r.text[:40]}'")
    print(f"P{i} (style={p.style.name}, align={p.alignment}): {runs_info if runs_info else '(empty)'}")

print("\n=== TABLES ===")
print(f"Jumlah tabel: {len(doc.tables)}")
for ti, t in enumerate(doc.tables):
    print(f"  Tabel {ti}: {len(t.rows)} rows x {len(t.columns)} cols")
    for row in t.rows[:3]:
        print("    ", [c.text.strip()[:30] for c in row.cells])

# section / margin
s = doc.sections[0]
print(f"\n=== SECTION ===")
print(f"  Page: {s.page_width} x {s.page_height}")
print(f"  Margins: top={s.top_margin} bottom={s.bottom_margin} left={s.left_margin} right={s.right_margin}")
