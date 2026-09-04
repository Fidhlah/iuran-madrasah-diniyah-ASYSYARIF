import docx
import json

path = r'C:\Users\Fidh\AppData\Local\hermes\tmp\Laporan Keuangan Bulan Juli 2026 (dgn KOP).docx'
doc = docx.Document(path)

ring = doc.tables[0]
lamp = doc.tables[1]

# ── Ambil tabel ringkasan ──
def ring_val(keyword):
    for row in ring.rows:
        cells = [c.text.strip() for c in row.cells]
        if cells[0] == keyword:
            return cells[1]
    return None

print("=== TABEL RINGKASAN (di file) ===")
for row in ring.rows:
    cells = [c.text.strip() for c in row.cells]
    print(f"  {cells[0]:25} {cells[1] if len(cells)>1 else ''}")

# ── Hitung ulang dari Lampiran ──
print("\n=== VERIFIKASI: hitung ulang dari Lampiran ===")
income_total = 0
expense_total = 0
income_count = 0
expense_count = 0
rows_data = []
for r in lamp.rows[1:]:  # skip header
    cells = [c.text.strip() for c in r.cells]
    if len(cells) < 5 or not cells[0].isdigit():
        continue
    no, tgl, jenis, ket, jml = cells[0], cells[1], cells[2], cells[3], cells[4]
    amt = int(jml)
    rows_data.append((no, tgl, jenis, ket, amt))
    if jenis == "Pemasukan":
        income_total += amt
        income_count += 1
    else:
        expense_total += amt
        expense_count += 1

print(f"  Lampiran: {len(rows_data)} transaksi")
print(f"    Pemasukan : {income_count} tx = Rp {income_total:,}")
print(f"    Pengeluaran: {expense_count} tx = Rp {expense_total:,}")

# ── Bandingkan dengan ringkasan ──
print("\n=== CEK COCOK TIDAK ===")
checks = [
    ("SPP Bulan Ini", None, "50000*51"),
    ("Total Pemasukan", income_total, None),
    ("PENGELUARAN RIIL", None, income_total),
]
# Total Pemasukan harus = income_total
tp = int(ring_val("Total Pemasukan").replace("Rp ","").replace(".",""))
pr = int(ring_val("PENGELUARAN RIIL").replace("Rp ","").replace(".",""))
alokasi = int(ring_val("ALOKASI KE KAS MDTA").replace("Rp ","").replace(".",""))

print(f"  Total Pemasukan file: Rp {tp:,}  | dari Lampiran: Rp {income_total:,}  → {'OK' if tp==income_total else 'MISMATCH'}")
# Pengeluaran RIIL = expense_total - alokasi
riil_hitung = expense_total - alokasi
print(f"  Pengeluaran RIIL file: Rp {pr:,}  | Expense lamp - alokasi: Rp {expense_total:,} - Rp {alokasi:,} = Rp {riil_hitung:,}  → {'OK' if pr==riil_hitung else 'MISMATCH'}")

# cek semua pemasukan income terdiri dari iuran 50k, pendaftaran
pendaftaran = sum(a for _,_,k,_,a in rows_data if 'pendaftaran' in k.lower())
spp = sum(a for _,_,k,_,a in rows_data if 'pemasukan' in k and 'iuran' in k.lower() and 'bulan juli' in k.lower())
tunggakan = sum(a for _,_,k,_,a in rows_data if 'pemasukan' in k and 'iuran' in k.lower() and 'bulan juli' not in k.lower())
print(f"  Rinci Pemasukan: SPP bini {spp:,} + tunggakan {tunggakan:,} + pendaftaran {pendaftaran:,} = {spp+tunggakan+pendaftaran:,} (harus {income_total:,})")

# Kas MDTA ada di lampiran?
mdta = [r for r in rows_data if 'kas mdta' in r[3].lower()]
print(f"\n  Transaksi 'Kas MDTA' di Lampiran: {len(mdta)} → {mdta}")

# duplikat No?
nos = [r[0] for r in rows_data]
print(f"  No urut: {len(nos)} baris, unik={len(set(nos))}")

# cek total saldo akhir
kb_akhir = int(ring_val("Kas Besar").replace("Rp ","").replace(".",""))
mdta_akhir = int(ring_val("Kas MDTA").replace("Rp ","").replace(".",""))
total_akhir = int(ring_val("TOTAL TUNAI AKHIR").replace("Rp ","").replace(".",""))
print(f"\n  Saldo akhir: KB {kb_akhir:,} + MDTA {mdta_akhir:,} = {kb_akhir+mdta_akhir:,} | TOTAL tunai file: {total_akhir:,} → {'OK' if kb_akhir+mdta_akhir==total_akhir else 'MISMATCH'}")