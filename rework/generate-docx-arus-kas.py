#!/usr/bin/env python3
"""Generate DOCX Laporan Keuangan Bulanan (Arus Kas) — format sama seperti Juni.

Template struktur (dari 'Laporan Keuangan Bulan Juni 2026.docx'):
  - Judul center bold, Times New Roman 12
  - Paragraf Lampiran/Perihal/Yth/Assalamu'alaikum
  - Tabel ringkasan (SALDO AWAL / PEMASUKAN / PENGELUARAN RIIL / ALOKASI / SALDO AKHIR)
  - CATATAN + Wassalamu'alaikum + tanda tangan
  - LAMPIRAN 1: RINCIAN TRANSAKSI KAS (DETAIL) — tabel No/Tanggal/Jenis/Keterangan/Jumlah
"""

import json
import sys
from datetime import datetime
from pathlib import Path

import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = Path(__file__).resolve().parent
DATA = SCRIPT_DIR / "isi-database-jan-jul.json"

# Konfigurasi bulan (bisa diubah untuk bulan lain)
BULAN = "Juli"
BULAN_UPPER = "JULI"
PERIODE = "JULI 2026"
DATE_PREFIX = "2026-07"
TANGGAL_DOK = "05 Juli 2026"
OUT_NAME = "Laporan Keuangan Bulan Juli 2026.docx"
OUT_DIR = Path(r"C:\Users\Fidh\AppData\Local\hermes\tmp")


def fmt_idr(n):
    """1234567 -> 1.234.567"""
    s = f"{n:,}".replace(",", ".")
    return s


def load_data():
    with open(DATA, encoding="utf-8") as f:
        db = json.load(f)
    return db["tables"]["finances"]["rows"]


def split_month(rows, prefix):
    return [r for r in rows if (r.get("date") or "")[:7] == prefix]


def compute_summary(rows):
    """Hitung ringkasan untuk satu bulan."""
    income = [r for r in rows if r["type"] == "income"]
    expense = [r for r in rows if r["type"] == "expense"]

    spp_bulan_ini = 0
    spp_tunggakan = 0
    pendaftaran = 0
    langsung_mdta = 0
    for r in income:
        amt = int(float(r["amount"]))
        d = r["description"].lower()
        if "kas mdta" in d or "suka rela" in d or "potongan tabungan" in d:
            langsung_mdta += amt
        elif "pendaftaran" in d:
            pendaftaran += amt
        elif "iuran" in d:
            if "bulan juli 2026" in d:
                spp_bulan_ini += amt
            else:
                spp_tunggakan += amt
        else:
            pendaftaran += amt  # fallback

    gaji = kas_mesjid = seragam = alokasi = belanja = 0
    for r in expense:
        amt = int(float(r["amount"]))
        d = r["description"].lower()
        if "honor" in d or "guru" in d:
            gaji += amt
        elif "mesjid" in d or "masjid" in d:
            kas_mesjid += amt
        elif "kas mdta" in d:
            alokasi += amt
        elif "seragam" in d:
            seragam += amt
        else:
            belanja += amt

    return {
        "spp_bulan_ini": spp_bulan_ini,
        "spp_tunggakan": spp_tunggakan,
        "pendaftaran": pendaftaran,
        "langsung_mdta": langsung_mdta,
        "total_pemasukan": sum(int(float(r["amount"])) for r in income),
        "gaji": gaji,
        "kas_mesjid": kas_mesjid,
        "seragam": seragam,
        "belanja": belanja,
        "total_pengeluaran_riil": gaji + kas_mesjid + seragam + belanja,
        "alokasi": alokasi,
    }


def compute_balance_until(rows, prefix):
    """Hitung saldo KB/MDTA akumulatif sampai dengan bulan prefix.

    Logika klasifikasi SAMA dengan generate-laporan-full.py (classify()):
    - expense 'guru/honor'  → KB out real
    - expense 'mesjid'      → KB out real
    - expense 'diambil dari (uang) kas mdta' → MDTA spend
    - expense 'kas mdta'    → transfer ke MDTA (alokasi)
    - expense 'seragam'     → MDTA spend
    - expense lainnya       → MDTA spend
    - income 'kas mdta' / 'suka rela' / 'potongan tabungan' → MDTA in
    - income lainnya        → KB in
    """
    fin = [r for r in rows if (r.get("date") or "")[:7] <= prefix]
    kb_in = kb_out_real = transfer_out = mdta_in = mdta_spend = 0
    for r in fin:
        amt = int(float(r["amount"]))
        d = r["description"].lower()
        t = r["type"]
        if t == "income":
            if "kas mdta" in d or "suka rela" in d or "potongan tabungan" in d:
                mdta_in += amt
            else:
                kb_in += amt
        else:
            guru = any(k in d for k in ["guru", "honor", "mengajar"])
            mesjid = any(k in d for k in ["kas masjid", "kas mesjid"])
            diambil = "diambil dari uang kas mdta" in d or "diambil dari kas mdta" in d
            kas_mdta = "kas mdta" in d

            if guru:
                kb_out_real += amt
            elif mesjid:
                kb_out_real += amt
            elif diambil:
                mdta_spend += amt
            elif kas_mdta:
                transfer_out += amt
            elif "seragam" in d:
                mdta_spend += amt
            else:
                mdta_spend += amt

    kb = kb_in - kb_out_real - transfer_out
    mdta = mdta_in + transfer_out - mdta_spend
    return kb, mdta


def main():
    rows = load_data()
    jul_rows = split_month(rows, DATE_PREFIX)

    # Ringkasan Juli
    s = compute_summary(jul_rows)

    # Saldo awal Juli = saldo akhir Juni (akumulasi sampai 2026-06)
    kb_awal, mdta_awal = compute_balance_until(rows, "2026-06")
    # Saldo akhir Juli
    kb_akhir = kb_awal + (s["spp_bulan_ini"] + s["spp_tunggakan"] + s["pendaftaran"] + s["langsung_mdta"]) - (s["gaji"] + s["kas_mesjid"]) - s["alokasi"]
    mdta_akhir = mdta_awal + s["alokasi"] + s["langsung_mdta"] - (s["belanja"] + s["seragam"])

    print(f"Saldo awal: KB {kb_awal:,} MDTA {mdta_awal:,} Total {kb_awal+mdta_awal:,}")
    print(f"Juli: masuk {s['total_pemasukan']:,} keluar riil {s['total_pengeluaran_riil']:,} alokasi {s['alokasi']:,}")
    print(f"Saldo akhir: KB {kb_akhir:,} MDTA {mdta_akhir:,} Total {kb_akhir+mdta_akhir:,}")

    # ── Build DOCX ──
    doc = docx.Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    def add_para(text="", bold=False, align=None, runs=None):
        p = doc.add_paragraph()
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if runs:
            for txt, b in runs:
                r = p.add_run(txt)
                r.bold = b
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
        else:
            r = p.add_run(text)
            r.bold = bold
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        return p

    # ── Halaman 1 ──
    add_para("LAPORAN KEUANGAN BULANAN (ARUS KAS)", bold=True, align="center")
    add_para(f"PERIODE {PERIODE}", bold=True, align="center")
    add_para()
    add_para("Lampiran\t: 1 (Satu) Berkas Rincian Transaksi", runs=[("Lampiran", True), ("\t: 1 (Satu) Berkas Rincian Transaksi ", False)])
    add_para("Perihal\t: Laporan Pertanggungjawaban Keuangan Bulanan", runs=[("Perihal", True), ("\t: Laporan Pertanggungjawaban Keuangan Bulanan ", False)])
    add_para("Yth. Kepala Madrasah Diniyah Asy Syarif ")
    add_para(" Di Tempat. ")
    add_para()
    add_para(f"Assalamu’alaikum Wr. Wb. Berikut adalah ringkasan arus kas Madrasah Diniyah Asy Syarif untuk periode {BULAN} 2026:")

    # ── Tabel Ringkasan (2 kolom) ──
    tbl = doc.add_table(rows=14, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows_data = [
        ("SALDO AWAL", f"Rp {fmt_idr(kb_awal + mdta_awal)}"),
        ("Kas Besar", f"Rp {fmt_idr(kb_awal)}"),
        ("Kas MDTA", f"Rp {fmt_idr(mdta_awal)}"),
        ("PEMASUKAN", ""),
        ("SPP Bulan Ini", f"Rp {fmt_idr(s['spp_bulan_ini'])}"),
        ("SPP Tunggakan", f"Rp {fmt_idr(s['spp_tunggakan'])}"),
        ("Uang Pendaftaran", f"Rp {fmt_idr(s['pendaftaran'])}"),
        ("Total Pemasukan", f"Rp {fmt_idr(s['total_pemasukan'])}"),
        ("PENGELUARAN RIIL", f"Rp {fmt_idr(s['total_pengeluaran_riil'])}"),
        ("ALOKASI KE KAS MDTA", f"Rp {fmt_idr(s['alokasi'])}"),
        ("SALDO AKHIR", ""),
        ("Kas Besar", f"Rp {fmt_idr(kb_akhir)}"),
        ("Kas MDTA", f"Rp {fmt_idr(mdta_akhir)}"),
        ("TOTAL TUNAI AKHIR", f"Rp {fmt_idr(kb_akhir + mdta_akhir)}"),
    ]
    for i, (k, v) in enumerate(rows_data):
        c0 = tbl.cell(i, 0)
        c1 = tbl.cell(i, 1)
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        p1 = c1.paragraphs[0]
        r0 = p0.add_run(k)
        r1 = p1.add_run(v)
        r0.font.name = "Times New Roman"
        r0.font.size = Pt(12)
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)
        # Bold untuk header + total
        if k in ("SALDO AWAL", "PEMASUKAN", "PENGELUARAN RIIL", "ALOKASI KE KAS MDTA", "SALDO AKHIR", "TOTAL TUNAI AKHIR"):
            r0.bold = True
        if k in ("SALDO AWAL", "TOTAL TUNAI AKHIR"):
            r1.bold = True

    add_para()
    add_para("CATATAN:", runs=[("CATATAN", True), (":", False)])
    add_para("Rincian setiap transaksi masuk dan keluar tersedia pada Lampiran 1. ")
    add_para()
    add_para("Wassalamu’alaikum Wr. Wb.")
    add_para()
    add_para("Bandung, " + TANGGAL_DOK, align="right")
    add_para("Heni Nuryati,", align="right")
    add_para()
    add_para()
    add_para()

    # ── Halaman 2: Lampiran ──
    doc.add_page_break()
    add_para("LAMPIRAN 1: RINCIAN TRANSAKSI KAS (DETAIL)", bold=True, align="center")
    add_para(PERIODE, bold=True, align="center")

    # Sort transaksi: date desc, dalam tanggal sama expense dulu
    def sort_key(r):
        d = r["date"][:10]
        prio = 0 if r["type"] == "expense" else 1
        return (d, prio, r["description"])

    tx_sorted = sorted(jul_rows, key=sort_key, reverse=True)

    lamp = doc.add_table(rows=1 + len(tx_sorted), cols=5)
    lamp.style = "Table Grid"
    headers = ["No", "Tanggal", "Jenis", "Keterangan", "Jumlah"]
    for ci, h in enumerate(headers):
        cell = lamp.cell(0, ci)
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(11)

    for ri, tx in enumerate(tx_sorted, start=1):
        dt = tx["date"][:10]
        # format d/m/yyyy (sama dgn Juni: 26/6/2026) → dari ISO 2026-07-31 = 31/7/2026
        y, m, d = dt.split("-")
        tgl = f"{int(d)}/{int(m)}/{y}"
        jenis = "Pemasukan" if tx["type"] == "income" else "Pengeluaran"
        ket = tx["description"].strip()
        jml = fmt_idr(int(float(tx["amount"])))
        vals = [str(ri), tgl, jenis, ket, jml]
        for ci, v in enumerate(vals):
            cell = lamp.cell(ri, ci)
            cell.text = ""
            r = cell.paragraphs[0].add_run(v)
            r.font.name = "Calibri"
            r.font.size = Pt(11)

    # Lebar kolom lampiran
    widths = [Cm(1.0), Cm(2.2), Cm(2.4), Cm(8.0), Cm(2.4)]
    for row in lamp.rows:
        for ci, w in enumerate(widths):
            row.cells[ci].width = w

    out_path = OUT_DIR / OUT_NAME
    doc.save(out_path)
    print(f"\n✅ DOCX selesai → {out_path}")
    print(f"   {len(tx_sorted)} transaksi di Lampiran")


if __name__ == "__main__":
    main()
