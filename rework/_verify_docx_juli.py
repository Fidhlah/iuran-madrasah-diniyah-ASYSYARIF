import docx

path = r'C:\Users\Fidh\AppData\Local\hermes\tmp\Laporan Keuangan Bulan Juli 2026.docx'
doc = docx.Document(path)

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"  {t[:90]}")

print("\n=== TABLE 0 (ringkasan) ===")
t0 = doc.tables[0]
for row in t0.rows:
    cells = [c.text.strip() for c in row.cells]
    print(f"  {cells}")

print("\n=== TABLE 1 (lampiran) ===")
t1 = doc.tables[1]
print(f"  Total rows: {len(t1.rows)} (header + {len(t1.rows)-1} transaksi)")
for ri, row in enumerate(t1.rows[:5]):
    cells = [c.text.strip()[:40] for c in row.cells]
    print(f"  R{ri}: {cells}")
print("  ...")
for ri, row in enumerate(t1.rows[-3:], start=len(t1.rows)-3):
    cells = [c.text.strip()[:40] for c in row.cells]
    print(f"  R{ri}: {cells}")
