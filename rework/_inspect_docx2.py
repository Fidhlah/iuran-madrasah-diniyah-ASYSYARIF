import docx
from docx.shared import Pt

path = r'D:/fidh/Asysyarif/Dokumen/Keuangan/Laporan Keuangan Bulanan/Juni 2026/Laporan Keuangan Bulan Juni 2026.docx'
doc = docx.Document(path)

# Inspect fonts of first few paragraphs
for i, p in enumerate(doc.paragraphs[:12]):
    runs_info = []
    for r in p.runs:
        fname = r.font.name
        fsize = r.font.size
        fbold = r.font.bold
        runs_info.append(f"[{fname} {fsize} bold={fbold}] '{r.text[:40]}'")
    print(f"P{i}: {runs_info}")

print()
print("=== TABLE 0 (ringkasan) detail ===")
t0 = doc.tables[0]
for ri, row in enumerate(t0.rows):
    cells = []
    for c in row.cells:
        txt = c.text.strip().replace('\n', ' | ')
        # check merged cells
        cells.append(txt)
    print(f"  R{ri}: {cells}")

print()
print("=== TABLE 1 (lampiran) header ===")
t1 = doc.tables[1]
for ri, row in enumerate(t1.rows[:4]):
    cells = [c.text.strip() for c in row.cells]
    print(f"  R{ri}: {cells}")

print()
print("=== TABLE 1 row 1 detail (col widths / fonts) ===")
r1 = t1.rows[1]
for ci, c in enumerate(r1.cells):
    for p in c.paragraphs:
        for run in p.runs:
            print(f"  col{ci}: font={run.font.name} size={run.font.size} bold={run.font.bold} text={run.text[:50]!r}")
