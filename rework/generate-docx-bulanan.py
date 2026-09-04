#!/usr/bin/env python3
"""Generate single DOCX: Laporan Keuangan Bulanan (sistem baru, fund separation)."""

import json, re
from datetime import datetime
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA = Path(__file__).parent / "isi-database.json"
OUT = Path(r"D:\fidh\Asysyarif\Dokumen\Keuangan\Laporan Keuangan Bulanan\Laporan Bulanan (Sistem Baru).docx")

# ── Load & dedup ──────────────────────────────────────────────
with open(DATA) as f:
    db = json.load(f)

finances = db["tables"]["finances"]["rows"]
seen = set()
uniq = []
for r in finances:
    if r["id"] not in seen:
        seen.add(r["id"])
        uniq.append(r)
finances = uniq

# ── Helpers ───────────────────────────────────────────────────
MONTH_NAMES = {1: "Januari", 2: "Februari", 3: "Maret", 4: "April",
               5: "Mei", 6: "Juni", 7: "Juli", 8: "Agustus",
               9: "September", 10: "Oktober", 11: "November", 12: "Desember"}
MONTH_NAMES_LOWER = {v.lower(): k for k, v in MONTH_NAMES.items()}

def parse_dt(s):
    try:
        return datetime.fromisoformat((s or "").replace("Z", ""))
    except Exception:
        return None

def rp(n):
    return f"Rp {n:,}".replace(",", ".")

# ── Classify income transaction (SPP type) ───────────────────
# Pattern: "... membayar iuran bulan [Month] [Year]"
RX_IURAN = re.compile(r"membayar iuran bulan (\w+) (\d{4})", re.IGNORECASE)

def classify_income(desc: str, tx_month: int, tx_year: int):
    dl = desc.lower()
    if "infaq" in dl or "shadaqah" in dl:
        return "Infaq & Lainnya"
    m = RX_IURAN.search(desc)
    if m:
        pay_month_name = m.group(1).lower()
        pay_year = int(m.group(2))
        pay_month = MONTH_NAMES_LOWER.get(pay_month_name)
        if pay_month is not None:
            if pay_year == tx_year and pay_month == tx_month:
                return "SPP Bulan Ini"
            elif pay_year < tx_year or (pay_year == tx_year and pay_month < tx_month):
                return "SPP Tunggakan"
            else:
                return "SPP Titipan"
    return "Infaq & Lainnya"


# ── Classify all transactions ──────────────────────────────────
classified = []
for tx in finances:
    dt = parse_dt(tx.get("date"))
    if not dt:
        continue
    desc = (tx.get("description") or "").lower()
    amt = int(tx["amount"])
    typ = tx.get("type", "")
    mk = dt.strftime("%Y-%m")
    tx_month = dt.month
    tx_year = dt.year

    if typ == "income":
        if "kas mdta" in desc:
            fund, cat, is_transfer, spp_type = "kas_mdta", "Pemasukan Langsung MDTA", False, None
        else:
            spp_t = classify_income(desc, tx_month, tx_year)
            fund, cat, is_transfer, spp_type = "kas_besar", spp_t, False, spp_t
    else:
        spp_type = None
        if "guru" in desc or "honor" in desc:
            fund, cat, is_transfer = "kas_besar", "Gaji & Honor Guru", False
        elif "kas mesjid" in desc or "kas masjid" in desc:
            fund, cat, is_transfer = "kas_besar", "Kas Mesjid", False
        elif "diambil dari uang kas mdta" in desc:
            fund, cat, is_transfer = "kas_mdta", "Belanja MDTA", False
        elif "kas mdta" in desc:
            fund, cat, is_transfer = "kas_besar", "Alokasi Kas MDTA", True
        elif "seragam" in desc:
            fund, cat, is_transfer = "kas_mdta", "Operasional (Seragam)", False
        else:
            fund, cat, is_transfer = "kas_mdta", "Belanja MDTA", False

    classified.append({
        "_dt": dt, "_month": mk,
        "fund": fund, "category": cat, "type": typ,
        "amount": amt, "is_transfer": is_transfer,
        "spp_type": spp_type,
    })

# ── Aggregate per month ───────────────────────────────────────
months = sorted(set(r["_month"] for r in classified))
month_data = {}

for m in months:
    rows = [r for r in classified if r["_month"] == m]
    d = {
        "kb_in": 0, "kb_out": 0, "kb_tr": 0, "md_in": 0, "md_sp": 0,
        "spp_bulan_ini": 0, "spp_tunggakan": 0, "spp_titipan": 0,
        "infaq": 0, "guru": 0, "mesjid": 0, "seragam_md": 0, "belanja_mdta": 0,
    }
    for r in rows:
        if r["is_transfer"]:
            d["kb_tr"] += r["amount"]
        elif r["fund"] == "kas_besar" and r["type"] == "income":
            d["kb_in"] += r["amount"]
            if r["spp_type"] == "SPP Bulan Ini":
                d["spp_bulan_ini"] += r["amount"]
            elif r["spp_type"] == "SPP Tunggakan":
                d["spp_tunggakan"] += r["amount"]
            elif r["spp_type"] == "SPP Titipan":
                d["spp_titipan"] += r["amount"]
            else:
                d["infaq"] += r["amount"]
        elif r["fund"] == "kas_besar" and r["type"] == "expense":
            d["kb_out"] += r["amount"]
            if "guru" in r["category"].lower():
                d["guru"] += r["amount"]
            elif "mesjid" in r["category"].lower():
                d["mesjid"] += r["amount"]
        elif r["fund"] == "kas_mdta" and r["type"] == "income":
            d["md_in"] += r["amount"]
        elif r["fund"] == "kas_mdta" and r["type"] == "expense":
            d["md_sp"] += r["amount"]
            if "seragam" in r["category"].lower():
                d["seragam_md"] += r["amount"]
            else:
                d["belanja_mdta"] += r["amount"]
    month_data[m] = d

# ── Helper: style a cell ──────────────────────────────────────
def style_cell(cell, text, bold=False, size=10, color=None, align=None, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if bg:
        shading = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg)
        shading.append(shd)


# ── Build DOCX ────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

GREEN = RGBColor(0x05, 0x96, 0x69)
RED = RGBColor(0xDC, 0x26, 0x26)
BLUE = RGBColor(0x25, 0x63, 0xEB)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GRAY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Title ─────────────────────────────────────────────────────
title = doc.add_heading("LAPORAN KEUANGAN BULANAN", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = DARK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Madrasah Diniyah Asy Syarif")
run.font.size = Pt(12)
run.bold = True
run.font.color.rgb = DARK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run("Periode Januari s.d. Juni 2026 · Sistem Baru (Fund Separation)")
run2.font.size = Pt(9)
run2.font.color.rgb = GRAY

doc.add_paragraph()  # spacer

# ── Color palette ────────────────────────────────────────────
HEADER_BG = "E2E8F0"
SUMMARY_BG = "F8FAFC"
SECTION_BG = "F1F5F9"

# ── Generate per-month table ──────────────────────────────────
kb_bal = 0
md_bal = 0

for m in months:
    d = month_data[m]
    month_num = int(m.split("-")[1])
    year = m.split("-")[0]
    month_name = MONTH_NAMES[month_num]

    kb_prev = kb_bal
    md_prev = md_bal

    kb_bal += d["kb_in"] - d["kb_out"] - d["kb_tr"]
    md_bal += d["kb_tr"] + d["md_in"] - d["md_sp"]

    total_inc = d["kb_in"] + d["md_in"]
    total_exp_real = d["kb_out"] + d["md_sp"]
    net = total_inc - total_exp_real
    total_awal = kb_prev + md_prev
    total_akhir = kb_bal + md_bal
    total_spp = d["spp_bulan_ini"] + d["spp_tunggakan"] + d["spp_titipan"] + d["infaq"]

    # ── Month heading ─────────────────────────────────────────
    doc.add_heading(f"{month_name} {year}", level=2)

    # ── Build rows for the main table ─────────────────────────
    # Each item: (label, value, bold, color, bg)
    items = []

    # Section: Saldo Awal
    items.append(("SALDO AWAL", rp(total_awal), True, DARK, SUMMARY_BG))
    items.append(("  Kas Besar", rp(kb_prev), False, BLUE, None))
    items.append(("  Kas MDTA", rp(md_prev), False, BLUE, None))
    items.append(("", "", False, None, None))  # spacer

    # Section: Pemasukan
    items.append(("PEMASUKAN", "", False, None, SECTION_BG))
    items.append(("  SPP Bulan Ini", rp(d["spp_bulan_ini"]), False, GREEN, None))
    if d["spp_tunggakan"]:
        items.append(("  SPP Tunggakan", rp(d["spp_tunggakan"]), False, GREEN, None))
    if d["spp_titipan"]:
        items.append(("  SPP Titipan", rp(d["spp_titipan"]), False, GREEN, None))
    if d["infaq"]:
        items.append(("  Infaq & Lainnya", rp(d["infaq"]), False, GREEN, None))
    if d["md_in"]:
        items.append(("  Pemasukan Langsung Kas MDTA", rp(d["md_in"]), False, GREEN, None))
    items.append(("  Total Pemasukan", rp(total_inc), True, GREEN, SUMMARY_BG))
    items.append(("", "", False, None, None))

    # Section: Pengeluaran Riil
    items.append(("PENGELUARAN RIIL", "", False, None, SECTION_BG))
    if d["guru"]:
        items.append(("  Gaji & Honor Guru", rp(d["guru"]), False, RED, None))
    if d["mesjid"]:
        items.append(("  Kas Mesjid", rp(d["mesjid"]), False, RED, None))
    if d["seragam_md"]:
        items.append(("  Operasional (Seragam)", rp(d["seragam_md"]), False, RED, None))
    if d["belanja_mdta"]:
        items.append(("  Belanja MDTA (ATK, print, dll)", rp(d["belanja_mdta"]), False, RED, None))
    items.append(("  Total Pengeluaran Riil", rp(total_exp_real), True, RED, SUMMARY_BG))
    items.append(("", "", False, None, None))

    # Section: Alokasi
    if d["kb_tr"]:
        items.append(("ALOKASI KE KAS MDTA", rp(d["kb_tr"]), True, BLUE, SECTION_BG))
        items.append(("", "", False, None, None))

    # Section: Saldo Akhir
    items.append(("SALDO AKHIR", "", False, None, SUMMARY_BG))
    items.append(("  Kas Besar", rp(kb_bal), True, BLUE, None))
    items.append(("  Kas MDTA", rp(md_bal), True, BLUE, None))
    items.append(("TOTAL TUNAI AKHIR", rp(total_akhir), True, GREEN, SUMMARY_BG))

    # ── Create table ──────────────────────────────────────────
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths
    for cell in table.columns[0].cells:
        cell.width = Cm(10)
    for cell in table.columns[1].cells:
        cell.width = Cm(5)

    for i, (label, value, bold, color, bg) in enumerate(items):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]

        # Label
        style_cell(c0, label, bold=bold, size=10, color=color if color else DARK, bg=bg)

        # Value (right-aligned)
        if value:
            align = WD_ALIGN_PARAGRAPH.RIGHT
            style_cell(c1, value, bold=bold, size=10 if not bold else 11,
                       color=color, align=align, bg=bg)

    doc.add_paragraph()  # spacer

# ── Save ──────────────────────────────────────────────────────
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"✅ Saved: {OUT}")
